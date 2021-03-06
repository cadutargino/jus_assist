import base64
import os
from unidecode import unidecode
import docx
import genderbr
import re
import streamlit as st
from PyPDF2 import PdfFileReader
import comarcas
from comarcas import grandes
from dict_crimes import dict_crimes, trafico, furto, lei_maria_penha, roubo, transito, med_prot, dict_med_prot, prisao, \
    dict_prisao, crimes_especie2, multa_penal_filename, blanck
from datetime import datetime
from num2words import num2words


def number_to_long_number(number_p):
    if number_p.find(',')!=-1:
        number_p = number_p.split(',')
        number_p1 = int(number_p[0].replace('.',''))
        number_p2 = int(number_p[1])
    else:
        number_p1 = int(number_p.replace('.',''))
        number_p2 = 0

    if number_p1 == 1:
        aux1 = ' real'
    else:
        aux1 = ' reais'

    if number_p2 == 1:
        aux2 = ' centavo'
    else:
        aux2 = ' centavos'

    text1 = ''
    if number_p1 > 0:
        text1 = num2words(number_p1,lang='pt_BR') + str(aux1)
    else:
        text1 = ''

    if number_p2 > 0:
        text2 = num2words(number_p2,lang='pt_BR') + str(aux2)
    else:
        text2 = ''

    if (number_p1 > 0 and number_p2 > 0):
        result = text1 + ' e ' + text2
    else:
        result = text1 + text2

    return result


def get_binary_file_downloader_html(bin_file, file_label='File'):
    with open(bin_file, 'rb') as f:
        data = f.read()
    bin_str = base64.b64encode(data).decode()
    href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{os.path.basename(bin_file)}">Download {file_label}</a>'
    return href


def read_pdf(file):
    pdfReader = PdfFileReader(file)
    count = pdfReader.numPages
    all_page_text = ""
    for i in range(count):
        page = pdfReader.getPage(i)
        all_page_text += page.extractText()

    return all_page_text


def extract_term(text, raw_reg_expression):
    Regex = re.compile(raw_reg_expression)
    mo = Regex.search(text)
    if mo is None:
        term = None
    else:
        term = mo.group()
    return term


def change_term_by_placeholder(d, parag, placeholder, term):
    """
     Insere termo na minuta pelo parágrafo e placeholder\n
     - d: docx.Document do pacote docx
     - parag: termo que identifica parágrafo
     - placeholder: termo existente no texto para identificá-lo no texto
     - term: termo a ser inserido na minuta
     """
    if term is not None:
        for para in range(len(d.paragraphs)):
            if parag in d.paragraphs[para].text:
                for i in range(len(d.paragraphs[para].runs)):
                    if placeholder in d.paragraphs[para].runs[i].text:
                        d.paragraphs[para].runs[i].text = term
                        d.paragraphs[para].runs[i].underline = False
    else:
        pass


def change_term_in_whole_document(d, placeholder, term, bold_und=None):
    """
    Insere termo na minuta pelo placeholder\n
    - d: docx.Document do pacote docx
    - placeholder: termo existente no texto para identificá-lo no texto
    - term: termo a ser inserido na minuta
    - bold_und: bold, underline, bold_underline, italic (formato do termo no texto final)
    """
    if term is not None:
        for para in range(len(d.paragraphs)):
            for run in range(len(d.paragraphs[para].runs)):
                if d.paragraphs[para].runs[run].text == placeholder:
                    d.paragraphs[para].runs[run].text = term
                    if bold_und == "bold":
                        d.paragraphs[para].runs[run].bold = True
                        d.paragraphs[para].runs[run].underline = False

                    elif bold_und == "underline":
                        d.paragraphs[para].runs[run].underline = True
                        d.paragraphs[para].runs[run].bold = False

                    elif bold_und == "bold_underline":
                        d.paragraphs[para].runs[run].underline = True
                        d.paragraphs[para].runs[run].bold = True
                    elif bold_und == "italic":
                        d.paragraphs[para].runs[run].italic = True
                        d.paragraphs[para].runs[run].underline = False
                        d.paragraphs[para].runs[run].bold = False
                    else:
                        d.paragraphs[para].runs[run].underline = False
                        d.paragraphs[para].runs[run].bold = False
    else:
        pass


def change_term_by_font_style(d, parag, style_match, term, case=None, bold_und=None):
    """
    Insere termo na minuta pelo parágrafo e estilo da fonte do placeholder\n
    - d: docx.Document do pacote docx
    - parag: termo existente no parágrafo para identificá-lo
    - style_match: underline, bold, italic (formato do placeholder)
    - term: termo a ser inserido na minuta
    - case: upper, lower (default = None)
    - bold_und: bold, underline, bold_underline, italic (formato do termo no texto final)
    """
    if term is not None:
        for para in range(len(d.paragraphs)):
            if parag in d.paragraphs[para].text:
                for i in range(len(d.paragraphs[para].runs)):
                    if style_match == "underline":
                        if d.paragraphs[para].runs[i].underline:
                            if case == "upper":
                                d.paragraphs[para].runs[i].text = term.upper()
                            else:
                                d.paragraphs[para].runs[i].text = term
                            d.paragraphs[para].runs[i].underline = False
                            if bold_und == "bold":
                                d.paragraphs[para].runs[i].bold = True
                                d.paragraphs[para].runs[i].underline = False

                            elif bold_und == "underline":
                                d.paragraphs[para].runs[i].underline = True
                                d.paragraphs[para].runs[i].bold = False

                            elif bold_und == "bold_underline":
                                d.paragraphs[para].runs[i].underline = True
                                d.paragraphs[para].runs[i].bold = True
                            elif bold_und == "italic":
                                d.paragraphs[para].runs[i].italic = True
                                d.paragraphs[para].runs[i].underline = False
                                d.paragraphs[para].runs[i].bold = False
                            else:
                                d.paragraphs[para].runs[i].underline = False
                                d.paragraphs[para].runs[i].bold = False
                    elif style_match == "bold":
                        if d.paragraphs[para].runs[i].bold:
                            if case == "upper":
                                d.paragraphs[para].runs[i].text = term.upper()
                            else:
                                d.paragraphs[para].runs[i].text = term
                            d.paragraphs[para].runs[i].bold = False
                            if bold_und == "bold":
                                d.paragraphs[para].runs[i].bold = True
                                d.paragraphs[para].runs[i].underline = False

                            elif bold_und == "underline":
                                d.paragraphs[para].runs[i].underline = True
                                d.paragraphs[para].runs[i].bold = False

                            elif bold_und == "bold_underline":
                                d.paragraphs[para].runs[i].underline = True
                                d.paragraphs[para].runs[i].bold = True
                            elif bold_und == "italic":
                                d.paragraphs[para].runs[i].italic = True
                                d.paragraphs[para].runs[i].underline = False
                                d.paragraphs[para].runs[i].bold = False
                            else:
                                d.paragraphs[para].runs[i].underline = False
                                d.paragraphs[para].runs[i].bold = False
                    elif style_match == "italic":
                        if d.paragraphs[para].runs[i].italic:
                            if case == "upper":
                                d.paragraphs[para].runs[i].text = term.upper()
                            else:
                                d.paragraphs[para].runs[i].text = term
                            d.paragraphs[para].runs[i].italic = False
                            if bold_und == "bold":
                                d.paragraphs[para].runs[i].bold = True
                                d.paragraphs[para].runs[i].underline = False

                            elif bold_und == "underline":
                                d.paragraphs[para].runs[i].underline = True
                                d.paragraphs[para].runs[i].bold = False

                            elif bold_und == "bold_underline":
                                d.paragraphs[para].runs[i].underline = True
                                d.paragraphs[para].runs[i].bold = True
                            elif bold_und == "italic":
                                d.paragraphs[para].runs[i].italic = True
                                d.paragraphs[para].runs[i].underline = False
                                d.paragraphs[para].runs[i].bold = False
                            else:
                                d.paragraphs[para].runs[i].underline = False
                                d.paragraphs[para].runs[i].bold = False
    else:
        pass


def titled_string_rectifier(string):
    list_string = string.split()
    new_list = []
    new_string = ""
    for word in list_string:
        if len(word) > 3 or len(list_string) < 3 or word not in list_string[1: -1]:
            new_list.append(word)
        else:
            new_list.append(word.lower())
    new_string = " ".join(new_list)
    return new_string


def translate_date(data):
    data = data.split(sep='/')
    meses = {'01': 'janeiro', '02': 'fevereiro', '03': 'março', '04': 'abril', '05': 'maio',
         '06': 'junho',
         '07': 'julho', '08': 'agosto', '09': 'setembro', '10': 'outubro', '11': 'novembro',
         '12': 'dezembro'}
    return data[0] + ' de ' + meses[data[1]] + ' de ' + data[2]


def transform_date(string_data):
    date_match = re.findall(r'\d{1,2}/\d{2}/\d{4}', string_data)
    datas_normal = []
    datas_extenso = []
    for date in date_match:
        datas_normal.append(date)
        datas_extenso.append(translate_date(date))
    for i in range(len(datas_normal)):
        string = re.sub(f'{datas_normal[i]}', f'{datas_extenso[i]}', string_data)
        string_data = string
    return string_data


def main():
    """Assistente de denúncia"""

    st.title('Assistente de redação de peças jurídicas')
    st.text('Minutas do MPSP a partir do PDF do boletim de ocorrência ou da certidão de multa penal.')
    nome_promotor = st.text_input("Insira o nome do(a) Promotor(a) de Justiça:")
    medida_protetiva = False
    prisao_flagrante = False
    denuncia = False
    multa_penal = False

    if st.checkbox("Denúncia"):
        denuncia = True
        crimes = ["Tráfico de drogas", "Furto (art. 155)", "Roubo (art. 157)", "Lei Maria da Penha", "Crimes de Trânsito"]
        genero_crime = st.selectbox("Tipo de crime", crimes)
        if genero_crime == "Tráfico de drogas":
            especie_crime = st.selectbox("Crime em espécie", trafico)
        elif genero_crime == "Furto (art. 155)":
            especie_crime = st.selectbox("Crime em espécie", furto)
        elif genero_crime == "Lei Maria da Penha":
            especie_crime = st.selectbox("Crime em espécie", lei_maria_penha)
        elif genero_crime == "Roubo (art. 157)":
            especie_crime = st.selectbox("Crime em espécie", roubo)
        elif genero_crime == "Crimes de Trânsito":
            especie_crime = st.selectbox("Crime em espécie", transito)
        nome_arquivo = dict_crimes[especie_crime]
        d = docx.Document(nome_arquivo)
    elif st.checkbox("Medidas Protetivas (Lei 11.340/2006)"):
        medida_protetiva = True
        especie_medida = st.selectbox("Parecer em medida protetiva de urgência", med_prot)
        nome_arquivo = dict_med_prot[especie_medida]
        d = docx.Document(nome_arquivo)
        # doc_file = st.file_uploader("Insira Boletim de Ocorrência", type=["pdf"])
        # if st.button("Process"):
        #     st.write("To be continued...")
    elif st.checkbox("Auto de prisão em flagrante - Parecer"):
        prisao_flagrante = True
        infracao = st.selectbox("Selecione o crime imputado ao indiciado ou selecione outros", crimes_especie2)
        if infracao == "Outro(s) crime(s)- selecione e insira no próximo campo":
            infracao = st.text_input("Insira os crimes aqui:")

        especie_prisao = st.selectbox("Parecer em auto de prisão em flagrante", prisao)
        nome_arquivo = dict_prisao[especie_prisao]
        d = docx.Document(nome_arquivo)
    elif st.checkbox("Execução de pena de multa"):
        multa_penal = True
        nome_arquivo = multa_penal_filename
        d = docx.Document(nome_arquivo)
        nome_comarca = st.text_input("Insira a comarca de endereçamento ou deixe em branco para usar a comarca de origem da sentença condenatória:")
    else:
        nome_arquivo = "em_branco"
        d = docx.Document(blanck)

    doc_file = st.file_uploader("Insira PDF do Boletim de Ocorrência ou da certidão de multa penal extraído do processo digital:", type=["pdf"])
    if st.button("Iniciar processamento"):
        if doc_file is not None and doc_file.type == "application/pdf":
            file_details = {"Filename": doc_file.name, "FileType": doc_file.type, "FileSize": doc_file.size}
            BOText = read_pdf(doc_file)

            # Extrai local
            Local = extract_term(BOText, r'Local: (.*?)CEP')
            # st.write(Local)
            if Local is not None and len(Local) < 120:
                try:
                    Local = Local[6:-5]
                    Local = Local.title().strip()
                    Local = titled_string_rectifier(Local)
                except:
                    Local = ""
            elif Local is not None:
                try:
                    Local = extract_term(BOText, r'Local: (.*?)Tipo')
                    city_match = re.findall(r'[^-]+-* *[sS][pP]', Local)
                    city = city_match[len(city_match)-1]
                    Local = Local.replace(city, "")
                    Local = Local[6:-5]
                    Local = Local.title().strip()
                    Local = titled_string_rectifier(Local)
                except:
                    Local = ""
            else:
                Local = ""
            # Extrai cidade
            try:
                cidade = extract_term(BOText, r'Local: (.*?)Tipo')
            except:
                cidade = ""
            if cidade is not None:
                try:
                    cidade2_match = re.findall(r'[^-]+-* *[sS][pP]', cidade)
                    cidade2 = cidade2_match[len(cidade2_match) - 1]
                    # cidade2 = extract_term(cidade, r'[^-]+-* *[sS][pP]')
                    cidade2 = cidade2[:-4].title().strip()
                    cidade2 = re.sub(r"\d+", "", cidade2)
                    cidade2 = titled_string_rectifier(cidade2)
                except:
                    cidade2 = ""
            else:
                cidade2 = ""



            comarc = ""
            for city in comarcas.comarcas2:
                if unidecode(cidade2) == unidecode(city):
                    comarc = comarcas.comarcas2[city]

            # Foro
            if multa_penal:
                if nome_comarca != "" and len(nome_comarca) > 2:
                    comarc = nome_comarca
                else:
                    cidade2 = extract_term(BOText, r'[Cc]omarca\/[Ff]oro.*Autor')
                    cidade2 = cidade2[12:-5]
                    for comarca in comarcas.comarcas2:
                        if comarca in cidade2:
                            cidade2 = comarca

                    comarc = comarcas.comarcas2[cidade2]

                processo = extract_term(BOText, r'\d{4,7}-\d{2}.\d{4}.\d.\d{2}.\d{4}')

                # Qualificação e nome:
                qualificacao = extract_term(BOText, r'[Dd]evedor.*?:.*?Local')
                qualificacao = qualificacao[51:-5]
                executado = qualificacao.split(",")[0]
                qualificacao = ", ".join(qualificacao.split(",")[1::])
                # Valor da multa normal e por extenso
                valor = extract_term(BOText, r'R\$.*,\d{2}')

                extenso = number_to_long_number(valor[2:])
                extenso = f"({extenso})"
                # Artigo
                artigo = extract_term(BOText, r'[Ii]nfraç..*?:.*?Senten')
                artigo = artigo[9:-6]
                if "ano" in artigo or "meses" in artigo:
                    artigo = f"a {artigo}"
                else:
                    artigo = f"como incurso no {artigo}"
            else:
                qualificacao =""
                extenso = ""
                artigo = ""
                executado = ""
                valor = ""


            if cidade2 == "S.Paulo" or cidade2 == "S. Paulo":
                cidade2 = "São Paulo"
            elif cidade2 == "S.Sebastiao" or cidade2 == "S. Sebastiao" or cidade2 == "S.Sebastião" or cidade2 == "S. Sebastião":
                cidade2 = "São Sebastião"

            if unidecode(cidade2) == unidecode(comarc):
                municipalidade = f"e comarca de {comarc}"
            else:
                municipalidade = f"de {cidade2} e comarca de {comarc}"

            if comarc in grandes:
                tipo_vara = "criminal"
            else:
                tipo_vara = "judicial"


            # Extrai indiciado/autor
            indiciado = ""
            indiciado = extract_term(BOText, r'(Indiciad([oa])|Autor|Investigad([oa])): -(.*?)-')
            if indiciado is not None and 'Autor' in indiciado:
                indiciado = indiciado[9:-2]
            elif indiciado is not None and "Investiga" in indiciado:
                indiciado = indiciado[14:-2]
            elif indiciado is not None:
                indiciado = indiciado[13:-2]
            else:
                indiciado = ""

            # Extrai vítima:
            vitima = extract_term(BOText, r'Vítima: -(.*?)-')
            if vitima is not None:
                if len(vitima) > 50:
                    vitima = vitima[10:50].title()
                    vitima = titled_string_rectifier(vitima)
                else:
                    vitima = vitima[10:-2].title()
                    vitima = titled_string_rectifier(vitima)
            else:
                vitima = ""

            # Extrai testemunha e condutor:
            testemunha = extract_term(BOText, r'Testemunha: -(.*?)-')
            if testemunha is not None:
                if len(testemunha) > 50:
                    testemunha = testemunha[13:50].title()
                    testemunha = titled_string_rectifier(testemunha)
                else:
                    testemunha = testemunha[13:-2].title()
                    testemunha = titled_string_rectifier(testemunha)
            else:
                testemunha = ""

            profissao_testemunha = extract_term(BOText, r'Testemunha: -(.*?)Profissão:(.*?)-')
            if profissao_testemunha is not None:
                profissao_testemunha = profissao_testemunha.lower()
                if "policial militar" in profissao_testemunha:
                    profissao_testemunha = "policial militar"
                elif "guarda civil" in profissao_testemunha:
                    profissao_testemunha = "guarda civil"
                elif "policial civil" in profissao_testemunha:
                    profissao_testemunha = "policial civil"
                elif 'policial rodoviario' in profissao_testemunha or "policial rodoviário" in profissao_testemunha:
                    profissao_testemunha = "policial rodoviário"
                elif "policial" not in profissao_testemunha or len(profissao_testemunha) > 25:
                    profissao_testemunha = ""
            else:
                profissao_testemunha = ""

            condutor = extract_term(BOText, r'Condutor: -(.*?)-')
            if condutor is not None:
                if len(condutor) > 50:
                    condutor = condutor[11:50].title()
                    condutor = titled_string_rectifier(condutor)
                else:
                    condutor = condutor[11:-2].title()
                    condutor = titled_string_rectifier(condutor)
            else:
                condutor = ""

            profissao_condutor = extract_term(BOText, r'Condutor: -(.*?)Profissão:(.*?)-')
            if profissao_condutor is not None:
                profissao_condutor = profissao_condutor.lower()
                if "policial militar" in profissao_condutor:
                    profissao_condutor = "policial militar"
                elif "guarda civil" in profissao_condutor:
                    profissao_condutor = "guarda civil"
                elif "policial civil" in profissao_condutor:
                    profissao_condutor = "policial civil"
                elif 'policial rodoviario' in profissao_condutor or "policial rodoviário" in profissao_condutor:
                    profissao_condutor = "policial rodoviário"
                elif "policial" not in profissao_condutor or len(profissao_condutor) > 25:
                    profissao_condutor = ""
            else:
                profissao_condutor = ""

            if profissao_condutor != "":
                condutor_profissao = f"{condutor}, {profissao_condutor}"
            else:
                condutor_profissao = f"{condutor}"

            if profissao_testemunha != "":
                testemunha_profissao = f"{testemunha}, {profissao_testemunha}"
            else:
                testemunha_profissao = f"{testemunha}"


            # Extrai data e hora
            DataHora = extract_term(BOText, r'Ocorrência: (.*?)Comu')
            # st.write(DataHora)
            if DataHora is not None and len(DataHora) < 90:
                DataHora = DataHora[:-4]
                DataHora = DataHora.lower()
                data_hora = DataHora[11:].strip()
                data_match = re.findall(r'\d{1,2}/\d{2}/\d{4}', data_hora)
                if data_match is not None and len(data_match) == 1:
                    data = data_match[0]
                    data_ext = translate_date(data)
                    # data = data.split(sep='/')
                    # meses = {'01': 'janeiro', '02': 'fevereiro', '03': 'março', '04': 'abril', '05': 'maio',
                    #          '06': 'junho',
                    #          '07': 'julho', '08': 'agosto', '09': 'setembro', '10': 'outubro', '11': 'novembro',
                    #          '12': 'dezembro'}
                    # data_ext = data[0] + ' de ' + meses[data[1]] + ' de ' + data[2]
                else:
                    data_ext = transform_date(data_hora)

                # data = DataHora.split()[1]
                hora_match = re.findall(r'\d{1,2}:\d{2}', data_hora)
                if hora_match is not None and len(hora_match) == 1:
                    hora = hora_match[0]
                else:
                    hora = ""


                if ":" in hora and len(hora) <= 5:
                    hora = hora.split(sep=':')
                    hora = 'por volta de ' + hora[0] + 'h' + hora[1] + 'min'
                else:
                    hora = ""
                    data_ext = transform_date(data_hora)
            else:
                hora = ""
                data_ext = ""


            # Insere data atual e grava na variável data_atual
            now = datetime.now()

            meses_2 = {1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril', 5: 'maio', 6: 'junho',
                       7: 'julho', 8: 'agosto', 9: 'setembro', 10: 'outubro', 11: 'novembro',
                       12: 'dezembro'}

            data_atual = f"{now.day} de {meses_2[now.month]} de {now.year}"

            # Extrai numero do processo
            if multa_penal:
                # Processo principal execução:
                numero = extract_term(BOText, r'\d{4,7}-\d{2}.\d{4}.\d.\d{2}.\d{4}')
            else:
                numRegex = re.compile(r'\d{4,7}-\d{2}.\d{4}.\d.\d{2}.\d{4}')
                mo_num = numRegex.search(doc_file.name)
                if mo_num is not None:
                    numero = mo_num.group()
                else:
                    numero = doc_file.name[:-4]



            # numero = doc_file.name[-29:-4]  # numero do processo para incluir no nome do arquivo final

            # Mostrando resultados
            # st.write(file_details)
            # st.write(BOText) # bom para checar o BO na íntegra
            st.subheader("Resumo dos dados extraídos")
            st.markdown(f"**local:** {Local}")
            st.markdown(f"**cidade:** {cidade2}")
            st.markdown(f"**comarca:** {comarc}")
            st.markdown(f"**executado:** {executado}")
            st.markdown(f"**qualificação:** {qualificacao}")
            st.markdown(f"**Valor da multa:** {valor}")
            st.markdown(f"**indiciado:** {indiciado}")
            st.markdown(f"**condutor:** {condutor_profissao}")
            st.markdown(f"**testemunha:** {testemunha_profissao}")
            st.markdown(f"**vítima**: {vitima}")
            st.markdown(f"**autos nº:** {numero}")
            st.markdown(f"**data do fato:** {data_ext}")
            st.markdown(f"**hora do fato:** {hora}")
            # st.markdown(f"**data atual:** {data_atual}")



            # Inserir nome do Promotor
            if nome_promotor is None or len(nome_promotor) <= 0:
                nome_promotor = None
                promotor_justica = "Promotor de Justiça"
            else:
                primeiro_nome_promotor = nome_promotor.split()[0]
                try:
                    sexo_promotor = genderbr.get_gender(primeiro_nome_promotor)
                    if sexo_promotor == "F":
                        sexo_promotor = "feminino"
                    else:
                        sexo_promotor = "masculino"
                    st.markdown(f"**Nome do(a) Promotor(a) de Justiça:** {nome_promotor}")
                    st.markdown(f"**Sexo do(a) Promotor(a) de Justiça:** {sexo_promotor}")
                    if sexo_promotor == "feminino":
                        promotor_justica = "Promotora de Justiça"
                    else:
                        promotor_justica = "Promotor de Justiça"
                except:
                    promotor_justiça = "Promotor de Justiça"

            change_term_in_whole_document(d, "Subscritor", nome_promotor, "bold")
            change_term_in_whole_document(d, "Promotor", promotor_justica, "bold")

            # Inserir a comarca na primeira linha e assinatura:
            change_term_by_font_style(d, "EXCELENTÍSSIMO", "underline", comarc, 'upper', "bold")
            change_term_by_font_style(d, "EXCELENTÍSSIMO", "italic", tipo_vara, 'upper', "bold")
            change_term_in_whole_document(d, "sede_do_juizo", comarc)

            # Inserir testemunhas:
            change_term_in_whole_document(d, "CONDUTOR1", condutor_profissao)
            change_term_in_whole_document(d, "TESTEMUNHA2", testemunha_profissao)
            change_term_in_whole_document(d, "vítima3", vitima)
            change_term_in_whole_document(d, "placeholder4", vitima)

            # inserir data atual (a atualização automatica do word não funciona direito)
            change_term_in_whole_document(d, "5TODAY5", data_atual)
            # Alterações DENUNCIA:
            if denuncia:
                # Trocar a data, cidade, comarca, endereço, horario e o nome do denunciado na Denúncia
                change_term_by_placeholder(d, "Consta", "data", data_ext)
                change_term_by_placeholder(d, "Consta", "municipalidade", municipalidade)
                change_term_by_placeholder(d, "Consta", "sede_do_juizo", comarc)
                change_term_by_placeholder(d, "Consta", "endereco", Local)
                change_term_by_placeholder(d, "Consta", "hora", hora)
                change_term_by_font_style(d, "Consta", "bold", indiciado, bold_und="bold")
                change_term_by_font_style(d, "Ante o exposto", "bold", indiciado, bold_und="bold")
                change_term_by_font_style(d, "Ofereço denúncia em separado", "bold", indiciado, bold_und="bold")
            else:
                pass

            # Alterações PARECER PRISAO
            if prisao_flagrante:
                # Trocar hora, endereço, cidade no Parecer em auto de prisão em flagrante:
                change_term_by_placeholder(d, "Trata-se de auto", "infracao", infracao)
                change_term_by_placeholder(d, "Trata-se de auto", "data", data_ext)
                change_term_by_placeholder(d, "Trata-se de auto", "municipalidade", cidade2)
                change_term_by_placeholder(d, "Trata-se de auto", "endereco", Local)
                change_term_by_placeholder(d, "Trata-se de auto", "hora", hora)
                change_term_by_placeholder(d, "Trata-se de auto", "indiciado4", indiciado)
            else:
                pass

            # Alterações MEDIDA PROTETIVA:
            if medida_protetiva:
                # Troca vítima na medida protetiva:
                change_term_by_placeholder(d, "Trata-se de expediente", "placeholder1", vitima)
                change_term_by_placeholder(d, "Trata-se de expediente", "indiciado4", indiciado)
            else:
                pass

            if multa_penal:
                change_term_by_placeholder(d, "respeitosamente, vem", "EXECUTADO", executado)
                change_term_by_placeholder(d, "respeitosamente, vem", "endereco", qualificacao)
                change_term_by_placeholder(d, "O executado foi condenado", "infracao", artigo)
                change_term_by_placeholder(d, "O valor da pena de multa aplicada", "montante", valor)
                change_term_by_placeholder(d, "O valor da pena de multa aplicada", "totalidade", extenso)
                change_term_by_placeholder(d, "Atribui-se", "montante", valor)
                change_term_by_placeholder(d, "Atribui-se", "totalidade", extenso)



            # Substituindo o número do processo no arquivo Word:
            for para in range(len(d.paragraphs)):
                for run in range(len(d.paragraphs[para].runs)):
                    d.paragraphs[para].runs[run].text = re.sub(r'\d{4,7}-\d{2}.\d{4}.\d.\d{2}.\d{4}', numero,
                                                               d.paragraphs[para].runs[run].text)

            d.save(f"{nome_arquivo[:-5]}_{numero}.docx")

            st.markdown(get_binary_file_downloader_html(f"{nome_arquivo[:-5]}_{numero}.docx", '  minuta da peça jurídica'),
                        unsafe_allow_html=True)

            os.remove(f"{nome_arquivo[:-5]}_{numero}.docx")
        else:
            st.warning('Insira um arquivo PDF válido e tente novamente.')


    # Aba sobre no sidebar:
    #
    # labels = ["Sobre o aplicativo", "Sobre o autor"]
    # choice = st.sidebar.selectbox("Sobre o aplicativo: / Sobre o autor:", labels)
    # if choice == "Sobre o aplicativo":
    st.markdown("---")
    st.subheader("Limitações e breves considerações:")
    st.markdown("* Destinado à confecção de peças simples, a partir de modelos para casos mais comuns.")
    st.markdown("* Extrai alguns dados do boletim de ocorrência ou da certidão de multa penal, tais como: data, horário e local do crime,"
                " cidade, comarca, nomes do indiciado, de testemunhas, vítima, número do processo, etc.")
    st.markdown("* Com esses dados, o aplicativo gera uma minuta da peça jurídica já com essas informações incorporados.")
    st.markdown('* O "recheio" da peça é apenas um modelo, sem modificação realizada pelo aplicativo. Na denúncia, a maior parte das'
                ' modificações é realizada no primeiro e no último parágrafos.')
    st.markdown("* O arquivo word da minuta fica disponível no link para download, após o processamento do PDF.")
    st.markdown("* O aplicativo não permanece com cópia do Boletim de Ocorrência, nem tampouco com qualquer informação extraída"
                " do boletim de ocorrência ou inserida pelo usuário.")
    st.markdown("* Por enquanto, extrai os dados de apenas um indiciado.")
    st.markdown("* O criador do aplicativo não é programador profissional. Pode haver falhas, confira o resultado final.")

        # elif choice == "Sobre o autor":
    st.markdown("---")
    st.subheader("Sobre o autor:")
    st.markdown("Este aplicativo foi desenvolvido por **Carlos Eduardo Targino da Silva**,"
                " 4º Promotor de Justiça de São Sebastião, para automação parcial da produção de peças jurídicas"
                " mais simples no âmbito do Ministério Público."
                " A ideia é tentar minimizar erros e evitar perda de tempo desnecessária com transcrição de dados.")
    st.markdown("Caso tenha alguma crítica ou sugestão, entre em contato por email: "
                "<mailto:carlostsilva@mpsp.mp.br>")
if __name__ == '__main__':
    main()
